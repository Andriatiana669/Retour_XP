# marches/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import MarcheDetaille, Carte
from django import forms
from django.core.exceptions import PermissionDenied
from django.views.decorators.http import require_GET
import io
import os
import docx
import docx2txt
from docx import Document
from django.http import HttpResponse
from docx import Document
from docx.shared import Cm
from django.conf import settings


@login_required
def description_view(request, marche_detaille_id):
    marche_detaille = get_object_or_404(MarcheDetaille, id=marche_detaille_id)
    cartes = marche_detaille.cartes.all()
    
    context = {
        'object': marche_detaille,
        'marche_detaille': marche_detaille,
        'cartes': cartes
    }
    return render(request, 'marches/description.html', context)


# ---------- FORMULAIRE RAPIDE CARTE ----------
class CarteForm(forms.ModelForm):
    class Meta:
        model = Carte
        fields = ['titre', 'objet', 'contenu', 'nom_client', 'date_retour', 'image']

# ---------- AJOUTER ----------
@login_required
def ajouter_carte_view(request, marche_detaille_id):
    marche = get_object_or_404(MarcheDetaille, id=marche_detaille_id)
    if not request.user.profile.est_responsable:
        raise PermissionDenied

    if request.method == 'POST':
        form = CarteForm(request.POST, request.FILES)
        if form.is_valid():
            carte = form.save(commit=False)
            carte.marche_detaille = marche
            carte.save()
            return redirect('marches:description', marche_detaille_id=marche.id)
    else:
        form = CarteForm()
    return render(request, 'marches/carte_form.html', {'form': form, 'marche': marche, 'titre': 'Ajouter'})

# ---------- MODIFIER ----------
@login_required
def modifier_carte_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)
    if not request.user.profile.est_responsable:
        raise PermissionDenied

    if request.method == 'POST':
        form = CarteForm(request.POST, request.FILES, instance=carte)
        if form.is_valid():
            form.save()
            return redirect('marches:description', marche_detaille_id=carte.marche_detaille.id)
    else:
        form = CarteForm(instance=carte)
    return render(request, 'marches/carte_form.html', {'form': form, 'marche': carte.marche_detaille, 'titre': 'Modifier'})

# ---------- SUPPRIMER ----------
@login_required
def supprimer_carte_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)
    if not request.user.profile.est_responsable:
        raise PermissionDenied

    if request.method == 'POST':
        marche_id = carte.marche_detaille.id
        carte.delete()
        return redirect('marches:description', marche_detaille_id=marche_id)
    return render(request, 'marches/carte_confirm_delete.html', {'carte': carte})

# ---------- VOIR (READ-ONLY) ----------
@login_required
def voir_carte_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)
    return render(request, 'marches/carte_detail.html', {'carte': carte})


@require_GET
def download_word_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)

    # création du document Word
    doc = Document()
    doc.add_heading(carte.titre, 0)
    doc.add_paragraph(f"Objet : {carte.objet or 'Non renseigné'}")
    doc.add_paragraph(f"Client : {carte.nom_client or 'Non renseigné'}")
    doc.add_paragraph(f"Date de retour : {carte.date_retour.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Dernière MAJ : {carte.date_maj.strftime('%d/%m/%Y %H:%M')}")
    doc.add_page_break()
    doc.add_heading('Contenu', level=1)
    doc.add_paragraph(carte.contenu)

    # sauvegarde en mémoire
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # réponse téléchargement
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{carte.titre}.docx"'
    return response