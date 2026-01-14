# marches/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import MarcheDetaille, Carte
from django import forms
from django.core.exceptions import PermissionDenied
from django.views.decorators.http import require_GET
import io
import os
import docx
import docx2txt
from docx import Document
from openpyxl import Workbook
import pandas as pd
from django.http import HttpResponse
from docx.shared import Cm
from django.conf import settings
import csv
from datetime import datetime
from django.utils import timezone
from openpyxl.styles import Alignment

from django.core.exceptions import PermissionDenied

def responsable_required(view_func):
    def wrapper(request, *args, **kwargs):
        profile = getattr(request.user, 'profile', None)
        if not profile or not profile.est_responsable:
            raise PermissionDenied
        return view_func(request, *args, **kwargs)
    return wrapper


@login_required
def description_view(request, marche_detaille_id):
    marche_detaille = get_object_or_404(MarcheDetaille, id=marche_detaille_id)
    cartes = marche_detaille.cartes.all()
    
    context = {
        'object': marche_detaille,
        'marche_detaille': marche_detaille,
        'cartes': cartes
    }
    return render(request, 'marches/description.html', context)


# ---------- FORMULAIRE RAPIDE CARTE ----------
class CarteForm(forms.ModelForm):
    class Meta:
        model = Carte
        # on AJOUTE les 3 champs
        fields = ['titre', 'objet', 'contenu', 'solution', 'date_retour', 'image',
                  'demarche_solution', 'nom_dossier', 'date_livraison']
        widgets = {
            'solution': forms.Select(attrs={'class': 'form-select'}),
            'demarche_solution': forms.Textarea(attrs={'rows': 3}),
            'date_livraison': forms.DateInput(attrs={'type': 'date'}),
        }
# ---------- AJOUTER ----------
@login_required
@responsable_required
def ajouter_carte_view(request, marche_detaille_id):
    marche = get_object_or_404(MarcheDetaille, id=marche_detaille_id)
    clients_marche = marche.nom_client.all()

    if request.method == 'POST':
        form = CarteForm(request.POST, request.FILES)
        clients_ids = request.POST.getlist('clients_choisis')

        # Vérification métier : au moins un client obligatoire
        if not clients_ids:
            form.add_error(None, "Vous devez sélectionner au moins un client.")

        if form.is_valid():
            carte = form.save(commit=False)
            carte.marche_detaille = marche

            # ForeignKey unique : on ne garde qu’un seul client
            clients_ids = request.POST.getlist('clients_choisis')
            if clients_ids:
                carte.nom_client_id = clients_ids[0]
            else:
                carte.nom_client = None

            # champs démarche
            carte.demarche_solution = request.POST.get('demarche_solution', '')
            carte.nom_dossier = request.POST.get('nom_dossier', '')
            date_liv = request.POST.get('date_livraison')
            carte.date_livraison = date_liv if date_liv else None

            carte.save()          # ← une seule fois
            return redirect('marches:description', marche_detaille_id=marche.id)
    else:
        form = CarteForm()

    return render(request, 'marches/carte_form.html', {
        'form': form,
        'marche': marche,
        'titre': 'Ajouter',
        'clients_marche': clients_marche,
    })

# ---------- MODIFIER ----------
@login_required
def modifier_carte_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)
    if not request.user.profile.est_responsable:
        raise PermissionDenied

    marche = carte.marche_detaille
    clients_marche = marche.nom_client.all()

    if request.method == 'POST':
        form = CarteForm(request.POST, request.FILES, instance=carte)
        clients_ids = request.POST.getlist('clients_choisis')

        # Vérification métier : au moins un client
        if not clients_ids:
            form.add_error(None, "Vous devez sélectionner au moins un client.")

        if form.is_valid():
            carte = form.save(commit=False)

            # ForeignKey unique : on garde le 1er choix
            if clients_ids:
                carte.nom_client_id = clients_ids[0]
            else:
                carte.nom_client = None

            # Champs démarche
            carte.demarche_solution = request.POST.get('demarche_solution', '')
            carte.nom_dossier       = request.POST.get('nom_dossier', '')
            date_liv = request.POST.get('date_livraison')
            carte.date_livraison    = date_liv if date_liv else None

            carte.save()
            return redirect('marches:description', marche_detaille_id=marche.id)
    else:
        form = CarteForm(instance=carte)

    return render(request, 'marches/carte_form.html', {
        'form': form,
        'marche': marche,
        'titre': 'Modifier',
        'clients_marche': clients_marche,
        # Pré-selection des clients déjà enregistrés (ForeignKey unique → 1 seul)
        'selected_client': carte.nom_client_id,
    })

# ---------- SUPPRIMER ----------
@login_required
def supprimer_carte_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)
    if not request.user.profile.est_responsable:
        raise PermissionDenied

    if request.method == 'POST':
        marche_id = carte.marche_detaille.id
        carte.delete()
        return redirect('marches:description', marche_detaille_id=marche_id)
    return render(request, 'marches/carte_confirm_delete.html', {'carte': carte})

# ---------- VOIR (READ-ONLY) ----------
@login_required
def voir_carte_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)
    return render(request, 'marches/carte_detail.html', {'carte': carte})


@require_GET
def download_word_view(request, carte_id):
    carte = get_object_or_404(Carte, id=carte_id)

    doc = Document()
    doc.add_heading(carte.titre, 0)

    # 1. Métadonnées
    doc.add_paragraph(f"Objet : {carte.objet or 'Non renseigné'}")
    doc.add_paragraph(f"Client : {carte.nom_client or 'Non renseigné'}")
    doc.add_paragraph(f"Date de retour : {carte.date_retour.strftime('%d/%m/%Y')}")

    # 2. Solution
    doc.add_paragraph(f"Solution : {carte.solution}")

    # 3. Bloc démarche (uniquement si la solution l’exige)
    if carte.solution.demande_demarche:
        doc.add_heading('Démarche & dossier', level=2)
        doc.add_paragraph(f"Démarche solution : {carte.demarche_solution or '–'}")
        doc.add_paragraph(f"Nom dossier : {carte.nom_dossier or '–'}")
        if carte.date_livraison:
            doc.add_paragraph(f"Date livraison : {carte.date_livraison.strftime('%d/%m/%Y')}")

    # 4. Contenu
    doc.add_page_break()
    doc.add_heading('Contenu', level=1)
    doc.add_paragraph(carte.contenu)

    # 5. Image (si présente)
    if carte.image and carte.image.file:
        image_path = carte.image.path
        if os.path.isfile(image_path):
            doc.add_picture(image_path, width=Cm(10))

    # 6. Footer : date de MAJ
    doc.add_paragraph()
    doc.add_paragraph(f"Dernière mise à jour : {carte.date_maj.strftime('%d/%m/%Y %H:%M')}")

    # Sauvegarde en mémoire
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{carte.titre}.docx"'
    return response


@login_required
@responsable_required
def export_cartes_xlsx(request, marche_detaille_id):
    marche = get_object_or_404(MarcheDetaille, id=marche_detaille_id)
    cartes = marche.cartes.select_related('solution', 'nom_client').all()

    # DataFrame
    df = pd.DataFrame([
        {
            'Titre': str(c.titre or ''),
            'Objet': str(c.objet or ''),
            'Client': str(c.nom_client or ''),
            'Date retour': c.date_retour,
            'Solution': str(c.solution or ''),
            'Demande démarche': 'Oui' if c.solution and c.solution.demande_demarche else 'Non',
            'Démarche solution': str(c.demarche_solution or ''),
            'Nom dossier': str(c.nom_dossier or ''),
            'Date livraison': pd.NaT if c.date_livraison is None else pd.Timestamp(c.date_livraison),
            'Date MAJ': pd.Timestamp(timezone.make_naive(c.date_maj) if c.date_maj.tzinfo else c.date_maj)
        }
        for c in cartes
    ])

    # Réponse HTTP
    filename = f"cartes_{marche.nom_marche}_{datetime.now():%Y%m%d_%H%M}.xlsx"
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    # Écriture directe dans la réponse
    df.to_excel(response, index=False, engine='openpyxl')
    return response


